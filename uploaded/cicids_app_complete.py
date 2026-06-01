#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
CICIDS2017 Intrusion Detection Desktop Application
Cybersecurity Network Analysis Tool with Gemini AI Integration
"""

import os
import sys
import threading
import time
import re
import numpy as np
import pandas as pd
import joblib
import customtkinter as ctk
from tkinter import filedialog, messagebox, ttk
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder, StandardScaler

try:
    from google import genai
    from google.genai import types
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("Warning: google.genai not installed. Run: pip install google-genai")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Run: pip install python-docx")

# ============================================================================
# CONFIGURATION
# ============================================================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, 'unified_balanced_60_40_cicids.csv')
MODEL_FILE = os.path.join(BASE_DIR, 'cyber_rf_model.joblib')
SCALER_FILE = os.path.join(BASE_DIR, 'cyber_scaler.joblib')
ENCODER_FILE = os.path.join(BASE_DIR, 'cyber_encoder.joblib')

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")  # set via env var; do not hardcode keys

GEMINI_SYSTEM_PROMPT = """You are an expert Senior Cyber Security Analyst & Digital Forensics Examiner. You will receive network traffic statistics classified by a Random Forest model. Your job is to draft a formal, highly professional, and comprehensive Executive & Technical Cyber Security Report fully in English. The report must strictly include the following clearly structured sections: 1. Executive Summary (overview of network health, total rows, and attack vs. benign ratio), 2. Detailed Attack Analysis (for each detected attack type, provide a scientific explanation of how it works, and map the top 5 network features to it, explaining scientifically why these metrics spiked, such as IAT or flow_duration), 3. Risk Assessment (severity and impact on infrastructure like Web, FTP, or SSH servers), 4. Mitigation & Remediation Strategies (clear, actionable engineering solutions for each attack, e.g., Rate Limiting, Timeout adjustments, Fail2ban, and CDNs). Maintain a formal, academic, and rigorous tone, avoiding generic descriptions."""

KEY_FEATURES = [
    'dst_port',
    'protocol',
    'flow_duration',
    'tot_fwd_pkts',
    'tot_bwd_pkts',
    'flow_byts_s',
]

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def load_training_data():
    """Load and preprocess training data from CICIDS CSV file."""
    if not os.path.exists(DATA_FILE):
        raise FileNotFoundError(f'Data file not found: {DATA_FILE}')

    df = pd.read_csv(DATA_FILE)
    
    # Clean data
    df.replace([np.inf, -np.inf], np.nan, inplace=True)
    df.dropna(axis=0, how='all', inplace=True)
    df.dropna(inplace=True)

    if 'label' not in df.columns:
        raise ValueError('Dataset must contain a "label" column')

    # Encode labels
    y = df['label'].astype(str)
    label_encoder = LabelEncoder()
    y_encoded = label_encoder.fit_transform(y)

    # Prepare features
    X = df.drop(columns=['label'])
    X_numeric = X.select_dtypes(include=[np.number]).copy()
    X_numeric.replace([np.inf, -np.inf], np.nan, inplace=True)
    X_numeric.fillna(X_numeric.mean(), inplace=True)

    feature_columns = X_numeric.columns.tolist()

    return X_numeric, y_encoded, label_encoder, feature_columns


def train_or_load_model():
    """Train or load pre-trained Random Forest model."""
    if os.path.exists(MODEL_FILE) and os.path.exists(SCALER_FILE) and os.path.exists(ENCODER_FILE):
        try:
            model = joblib.load(MODEL_FILE)
            scaler = joblib.load(SCALER_FILE)
            label_encoder = joblib.load(ENCODER_FILE)
            feature_columns = model.feature_columns
            return model, scaler, label_encoder, feature_columns
        except Exception as e:
            print(f"Failed to load saved model: {e}. Retraining...")

    # Train new model
    X, y, label_encoder, feature_columns = load_training_data()
    
    # Scale features
    scaler = StandardScaler()
    X_scaled = scaler.fit_transform(X)
    
    # Split data
    X_train, X_test, y_train, y_test = train_test_split(
        X_scaled, y, test_size=0.2, stratify=y, random_state=42
    )

    # Train Random Forest
    model = RandomForestClassifier(
        n_estimators=200,
        class_weight='balanced',
        random_state=42,
        n_jobs=-1,
    )
    model.fit(X_train, y_train)
    model.feature_columns = feature_columns

    # Save model
    joblib.dump(model, MODEL_FILE)
    joblib.dump(scaler, SCALER_FILE)
    joblib.dump(label_encoder, ENCODER_FILE)

    print(f"Model trained and saved. Accuracy: {model.score(X_test, y_test):.4f}")
    return model, scaler, label_encoder, feature_columns


def call_gemini_api(statistical_report: str) -> str:
    """Send statistical report to Gemini API and get security analysis using the modern SDK."""
    if not GEMINI_AVAILABLE:
        raise RuntimeError('google.genai not installed. Run: pip install google-genai')

    try:
        # تأسيس العميل بالطريقة الصحيحة للـ SDK الحديث
        client = genai.Client(api_key=GEMINI_API_KEY)

        prompt = f"""{GEMINI_SYSTEM_PROMPT}

====================================
إحصائيات تحليل حركة الشبكة:
====================================

{statistical_report}

====================================
الآن، اكتب التقرير الأمني الشامل:
====================================
"""

        # Official and stable method for model invocation with system instruction injection
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                system_instruction=GEMINI_SYSTEM_PROMPT,
                temperature=0.2,
            ),
        )

        # استخراج النص المرتجع بالطريقة السليمة والمباشرة للـ SDK الحديث
        if response.text:
            return response.text
        else:
            raise RuntimeError('لم يقم نموذج Gemini بإرجاع أي نصوص للتقرير.')

    except Exception as e:
        raise RuntimeError(f'Gemini API Error: {str(e)}')


def markdown_to_docx(markdown_text: str, output_path: str) -> None:
    """Convert Markdown text to professionally formatted DOCX."""
    if not DOCX_AVAILABLE:
        raise RuntimeError('python-docx not installed. Run: pip install python-docx')
    
    doc = Document()
    doc.add_heading('Digital Forensics & Network Traffic Analysis Report (CICIDS2017)', level=0)
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.rstrip()
        
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Handle headings
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue
        elif line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=2)
            continue
        elif line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=3)
            continue
        
        # Handle bullet points
        if line.startswith('- '):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            continue
        elif line.startswith('* '):
            text = line[2:].strip()
            doc.add_paragraph(text, style='List Bullet')
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.+)', line)
            if match:
                text = match.group(2).strip()
                doc.add_paragraph(text, style='List Number')
                continue
        
        # Regular paragraph
        doc.add_paragraph(line)
    
    doc.save(output_path)


# ============================================================================
# MAIN APPLICATION CLASS
# ============================================================================

class CICIDSApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        # Window setup
        self.title('CICIDS2017 Intrusion Detection System')
        self.geometry('1100x750')
        self.minsize(1000, 700)
        
        ctk.set_appearance_mode('dark')
        ctk.set_default_color_theme('blue')
        
        # Initialize model
        try:
            self.model, self.scaler, self.label_encoder, self.feature_columns = train_or_load_model()
        except Exception as e:
            messagebox.showerror('Model Error', f'Failed to load/train model:\n{str(e)}')
            self.destroy()
            return
        
        # Create tabs
        self.tabview = ctk.CTkTabview(self)
        self.tabview.pack(padx=20, pady=20, expand=True, fill='both')
        
        self.tabview.add('Manual Prediction')
        self.tabview.add('Bulk Analysis')
        
        self.create_manual_tab(self.tabview.tab('Manual Prediction'))
        self.create_bulk_tab(self.tabview.tab('Bulk Analysis'))
    
    def create_manual_tab(self, parent):
        """Create manual prediction interface."""
        main_frame = ctk.CTkFrame(parent)
        main_frame.pack(expand=True, fill='both', padx=15, pady=15)
        
        # Title
        title = ctk.CTkLabel(
            main_frame,
            text='Manual Network Traffic Prediction',
            font=ctk.CTkFont(size=16, weight='bold')
        )
        title.pack(pady=(0, 20))
        
        # Input fields
        input_frame = ctk.CTkFrame(main_frame)
        input_frame.pack(side='left', fill='both', padx=(0, 20), expand=True)
        
        input_title = ctk.CTkLabel(
            input_frame,
            text='Feature Values:',
            font=ctk.CTkFont(size=12, weight='bold')
        )
        input_title.pack(anchor='w', pady=(0, 10))
        
        self.manual_entries = {}
        for feature in KEY_FEATURES:
            frame = ctk.CTkFrame(input_frame)
            frame.pack(fill='x', pady=5)
            
            label = ctk.CTkLabel(frame, text=feature + ':', width=120)
            label.pack(side='left', padx=(0, 10))
            
            entry = ctk.CTkEntry(frame, width=200)
            entry.pack(side='left', fill='x', expand=True)
            self.manual_entries[feature] = entry
        
        # Predict button
        predict_btn = ctk.CTkButton(
            input_frame,
            text='Predict Traffic Type',
            command=self.manual_predict,
            height=40,
            font=ctk.CTkFont(size=12, weight='bold')
        )
        predict_btn.pack(fill='x', pady=(20, 10))
        
        # Output
        output_label = ctk.CTkLabel(
            input_frame,
            text='Prediction Result:',
            font=ctk.CTkFont(size=11, weight='bold')
        )
        output_label.pack(anchor='w', pady=(10, 5))
        
        self.manual_output = ctk.CTkTextbox(input_frame, height=150)
        self.manual_output.pack(fill='both', expand=True)
        
        # Info panel
        info_frame = ctk.CTkFrame(main_frame)
        info_frame.pack(side='right', fill='both', expand=True)
        
        info_title = ctk.CTkLabel(
            info_frame,
            text='Model Information',
            font=ctk.CTkFont(size=12, weight='bold')
        )
        info_title.pack(anchor='w', pady=(0, 10))
        
        info_text = f"""Model Type: Random Forest Classifier
Features: {len(self.feature_columns)}
Classes: {len(self.label_encoder.classes_)}

Attack Types:
{chr(10).join([f'  • {label}' for label in self.label_encoder.classes_])}
"""
        
        info_textbox = ctk.CTkTextbox(info_frame, height=300)
        info_textbox.pack(fill='both', expand=True)
        info_textbox.insert('0.0', info_text)
        info_textbox.configure(state='disabled')
    
    def manual_predict(self):
        """Handle manual prediction."""
        try:
            # Prepare input vector
            input_data = {}
            for feature, entry in self.manual_entries.items():
                value_str = entry.get().strip()
                if value_str:
                    input_data[feature] = float(value_str)
                else:
                    input_data[feature] = 0.0
            
            # Create feature vector matching model's feature order
            input_vector = np.array([input_data.get(f, 0.0) for f in self.feature_columns]).reshape(1, -1)
            
            # Scale
            input_scaled = self.scaler.transform(input_vector)
            
            # Predict
            prediction_encoded = self.model.predict(input_scaled)
            prediction_proba = self.model.predict_proba(input_scaled)[0]
            prediction_text = self.label_encoder.inverse_transform(prediction_encoded)[0]
            confidence = np.max(prediction_proba)
            
            # Display result
            result = f"Prediction: {prediction_text}\nConfidence: {confidence:.2%}"
            
            self.manual_output.configure(state='normal')
            self.manual_output.delete('0.0', 'end')
            self.manual_output.insert('0.0', result)
            self.manual_output.configure(state='disabled')
        except Exception as e:
            messagebox.showerror('Prediction Error', f'Error during prediction:\n{str(e)}')
    
    def create_bulk_tab(self, parent):
        """Create bulk analysis interface."""
        main_frame = ctk.CTkFrame(parent)
        main_frame.pack(expand=True, fill='both', padx=15, pady=15)
        
        # Title
        title = ctk.CTkLabel(
            main_frame,
            text='Bulk Network Analysis & AI Security Report',
            font=ctk.CTkFont(size=16, weight='bold')
        )
        title.pack(pady=(0, 15))
        
        # Button frame
        btn_frame = ctk.CTkFrame(main_frame)
        btn_frame.pack(fill='x', pady=(0, 15))
        
        upload_btn = ctk.CTkButton(
            btn_frame,
            text='📁 Choose CSV File',
            command=self.choose_file,
            height=40,
            font=ctk.CTkFont(size=12, weight='bold')
        )
        upload_btn.pack(side='left', padx=(0, 10))
        
        analyze_btn = ctk.CTkButton(
            btn_frame,
            text='▶ Start Analysis & Generate Report',
            command=self.start_analysis,
            height=40,
            font=ctk.CTkFont(size=12, weight='bold')
        )
        analyze_btn.pack(side='left', padx=(0, 10))
        
        self.file_label = ctk.CTkLabel(btn_frame, text='No file selected', text_color='gray')
        self.file_label.pack(side='left', fill='x', expand=True)
        
        # Progress bar
        self.progress_bar = ctk.CTkProgressBar(main_frame)
        self.progress_bar.pack(fill='x', pady=(0, 10))
        self.progress_bar.set(0)
        
        # Status label
        self.status_label = ctk.CTkLabel(main_frame, text='Ready', text_color='gray')
        self.status_label.pack(anchor='w', pady=(0, 10))
        
        # Output textbox
        self.bulk_output = ctk.CTkTextbox(main_frame)
        self.bulk_output.pack(fill='both', expand=True)
        
        self.selected_file = None
    
    def choose_file(self):
        """Choose CSV or Excel file for analysis."""
        file_path = filedialog.askopenfilename(
            title='Select Network Traffic File (CSV or Excel)',
            filetypes=[
                ('Supported Files', '*.csv *.xlsx *.xls'),
                ('CSV Files', '*.csv'),
                ('Excel Files', '*.xlsx *.xls'),
                ('All Files', '*.*'),
            ]
        )
        if file_path:
            self.selected_file = file_path
            self.file_label.configure(text=f'File: {os.path.basename(file_path)}')
    
    def start_analysis(self):
        """Start bulk analysis with Gemini report generation."""
        if not self.selected_file:
            messagebox.showwarning('No File', 'Please select a CSV or Excel file first.')
            return
        
        # Run analysis in background thread
        thread = threading.Thread(target=self.run_analysis, daemon=True)
        thread.start()
    
    def run_analysis(self):
        """Run the analysis pipeline."""
        try:
            self.progress_bar.set(0)
            self.status_label.configure(text='Loading file...')
            self.update()
            
            # Load file
            if self.selected_file.lower().endswith('.csv'):
                df = pd.read_csv(self.selected_file)
            else:
                df = pd.read_excel(self.selected_file)
            
            if df.empty:
                raise ValueError('File is empty')
            
            # Clean data
            self.status_label.configure(text='Cleaning data...')
            self.update()
            
            df.replace([np.inf, -np.inf], np.nan, inplace=True)
            df.fillna(df.mean(numeric_only=True), inplace=True)
            
            # Remove label column if exists
            if 'label' in df.columns:
                df = df.drop(columns=['label'])
            
            # Ensure all required columns exist and reorder
            missing_cols = [col for col in self.feature_columns if col not in df.columns]
            if missing_cols:
                raise ValueError(f'Missing required columns: {", ".join(missing_cols)}')
            
            df_clean = df[self.feature_columns].copy()
            
            # Make predictions
            self.status_label.configure(text='Making predictions...')
            self.update()
            self.progress_bar.set(0.3)
            
            X_scaled = self.scaler.transform(df_clean)
            predictions_encoded = self.model.predict(X_scaled)
            predictions = self.label_encoder.inverse_transform(predictions_encoded)
            
            # Generate statistics
            self.status_label.configure(text='Generating statistics...')
            self.update()
            self.progress_bar.set(0.5)
            
            report_text = self.generate_statistics(df_clean, predictions)
            
            # Display in UI
            self.bulk_output.configure(state='normal')
            self.bulk_output.delete('0.0', 'end')
            self.bulk_output.insert('0.0', report_text)
            self.bulk_output.configure(state='disabled')
            
            # Call Gemini API
            self.status_label.configure(text='Calling Gemini API...')
            self.update()
            self.progress_bar.set(0.7)
            
            if not GEMINI_AVAILABLE:
                raise RuntimeError('Google GenAI not available. Install: pip install google-genai')
            
            gemini_response = call_gemini_api(report_text)
            
            # Convert to DOCX
            self.status_label.configure(text='Converting to Word format...')
            self.update()
            self.progress_bar.set(0.85)
            
            if not DOCX_AVAILABLE:
                raise RuntimeError('python-docx not available. Install: pip install python-docx')
            
            # Ask for save location
            save_path = filedialog.asksaveasfilename(
                title='Save Security Analysis Report',
                defaultextension='.docx',
                filetypes=[('Word Documents', '*.docx')]
            )
            
            if not save_path:
                self.status_label.configure(text='Operation cancelled.')
                self.progress_bar.set(0)
                return
            
            markdown_to_docx(gemini_response, save_path)
            
            self.status_label.configure(text=f'✓ Report saved: {os.path.basename(save_path)}')
            self.progress_bar.set(1.0)
            
            messagebox.showinfo('Success', f'Security analysis report generated successfully!\n\nSaved to:\n{save_path}')
            
        except Exception as e:
            self.status_label.configure(text='Error occurred')
            self.progress_bar.set(0)
            messagebox.showerror('Analysis Error', f'Error during analysis:\n{str(e)}')
    
    def generate_statistics(self, X_df: pd.DataFrame, predictions: np.ndarray) -> str:
        """Generate statistical report from predictions."""
        total_rows = len(predictions)
        
        # Count predictions
        unique_classes, counts = np.unique(predictions, return_counts=True)
        class_counts = dict(zip(unique_classes, counts))
        
        benign_count = class_counts.get('Benign', 0)
        attack_count = total_rows - benign_count
        
        benign_indexes = np.where(predictions == 'Benign')[0].tolist()
        attack_indexes = np.where(predictions != 'Benign')[0].tolist()
        
        # Top features
        importances = pd.Series(self.model.feature_importances_, index=self.feature_columns)
        top_features = importances.nlargest(5)
        
        # Generate report text
        report_lines = [
            '=' * 60,
            'NETWORK TRAFFIC ANALYSIS REPORT',
            'Random Forest Classification Results',
            '=' * 60,
            '',
            f'Total Rows Analyzed: {total_rows}',
            f'Benign Traffic: {benign_count} ({benign_count/total_rows*100:.1f}%)',
            f'Attack Traffic: {attack_count} ({attack_count/total_rows*100:.1f}%)',
            '',
            'Benign Row Indexes (1-based):',
            f'{", ".join(str(i+1) for i in benign_indexes[:50])}{"..." if len(benign_indexes) > 50 else ""}',
            '',
            'Attack Row Indexes (1-based):',
            f'{", ".join(str(i+1) for i in attack_indexes[:50])}{"..." if len(attack_indexes) > 50 else ""}',
            '',
            'Attack Type Distribution:',
        ]
        
        for attack_type, count in sorted(class_counts.items()):
            if attack_type != 'Benign':
                report_lines.append(f'  {attack_type}: {count} rows ({count/total_rows*100:.1f}%)')
        
        report_lines.extend([
            '',
            'Top 5 Most Important Features:',
        ])
        
        for feat, importance in top_features.items():
            report_lines.append(f'  {feat}: {importance:.4f}')
        
        report_lines.extend([
            '',
            'Statistical Summary:',
            f'Mean flow duration: {X_df["flow_duration"].mean():.2f}',
            f'Mean forward packets: {X_df.get("tot_fwd_pkts", pd.Series([0])).mean():.2f}',
            f'Mean backward packets: {X_df.get("tot_bwd_pkts", pd.Series([0])).mean():.2f}',
            f'Mean bytes/sec: {X_df.get("flow_byts_s", pd.Series([0])).mean():.2f}',
            '',
        ])
        
        return '\n'.join(report_lines)


# ============================================================================
# MAIN EXECUTION
# ============================================================================

if __name__ == '__main__':
    app = CICIDSApp()
    app.mainloop()
