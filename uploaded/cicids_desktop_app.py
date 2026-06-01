import os
import sys
import traceback
import threading
import time
import re
import numpy as np
import pandas as pd
import joblib
import customtkinter as ctk
from tkinter import filedialog, messagebox
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.preprocessing import LabelEncoder

try:
    import google.generativeai as genai
    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ملاحظة: تأكد من تثبيت المكتبات التالية قبل التشغيل:
# pip install customtkinter scikit-learn pandas joblib openpyxl google-generativeai python-docx

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_FILE = os.path.join(BASE_DIR, 'unified_balanced_60_40_cicids.csv')
MODEL_FILE = os.path.join(BASE_DIR, 'cicids_rf_model.joblib')

MANUAL_FEATURES = [
    ('Destination Port', 'dst_port'),
    ('Flow Duration', 'flow_duration'),
    ('Total Fwd Packets', 'tot_fwd_pkts'),
    ('Total Backward Packets', 'tot_bwd_pkts'),
    ('Fwd Packet Length Max', 'fwd_pkt_len_max'),
    ('Flow Bytes/s', 'flow_byts_s'),
]

FEATURE_DESCRIPTIONS = {
    'dst_port': 'يشير إلى المنفذ المستهدف، والذي قد يظهر نمط هجوم معين عند تغييره أو تكراره.',
    'flow_duration': 'مدة التدفق تساعد على كشف جلسات حركة غير معتادة أو طويلة جداً.',
    'tot_fwd_pkts': 'عدد الحزم الأمامية يدل على كثافة الاتصال واحتمالية هجوم مثل DDoS.',
    'tot_bwd_pkts': 'عدد الحزم العكسية يكشف التوازن بين المرسل والمستقبل وحركة غير طبيعية.',
    'fwd_pkt_len_max': 'أقصى طول حزمة أمامية يمكن أن يكشف عن حزم هجوم كبيرة أو غير متوقعة.',
    'flow_byts_s': 'معدل البايت في الثانية يحدد شدة التدفق ويستخدم في التفريق بين هجوم وترافيك طبيعي.',
}

# English descriptions for report output
FEATURE_DESCRIPTIONS_EN = {
    'dst_port': 'Destination port: indicates the targeted service or host port, useful to detect scanning and targeted attacks.',
    'flow_duration': 'Flow duration: session length which can indicate long-running suspicious connections.',
    'tot_fwd_pkts': 'Total forward packets: measures request intensity from the source towards the destination.',
    'tot_bwd_pkts': 'Total backward packets: measures responses from the destination to the source.',
    'fwd_pkt_len_max': 'Forward packet max length: large unusual packets can be a sign of certain attacks.',
    'flow_byts_s': 'Flow bytes/s: bytes per second indicating the volume and intensity of traffic.',
}

DEFAULT_GEMINI_SYSTEM_PROMPT = """You are an expert cybersecurity analyst. You will receive a network traffic analysis report and must provide a professional, detailed security analysis report based on the data. Structure your response in clear sections: Executive Summary, Key Findings, Attack Analysis, Recommendations, and Conclusion. Use professional security terminology and make actionable recommendations."""


def load_training_data():
    if not os.path.exists(DATA_FILE):
        raise FileNotFoundError(f'لم يتم العثور على ملف البيانات: {DATA_FILE}')

    df = pd.read_csv(DATA_FILE)
    df.replace([np.inf, -np.inf], np.nan, inplace=True)
    df.dropna(axis=0, how='all', inplace=True)
    df.dropna(inplace=True)

    if 'label' not in df.columns:
        raise ValueError('الملف لا يحتوي على عمود label')

    y = df['label'].astype(str)
    label_encoder = LabelEncoder()
    y_encoded = label_encoder.fit_transform(y)

    X = df.drop(columns=['label'])
    X_numeric = X.select_dtypes(include=[np.number]).copy()

    X_numeric.replace([np.inf, -np.inf], np.nan, inplace=True)
    X_numeric.fillna(X_numeric.mean(), inplace=True)

    feature_means = X_numeric.mean()
    feature_columns = X_numeric.columns.tolist()

    return X_numeric, y_encoded, label_encoder, feature_means, feature_columns


def train_or_load_model():
    if os.path.exists(MODEL_FILE):
        try:
            loaded = joblib.load(MODEL_FILE)
            model = loaded['model']
            label_encoder = loaded['label_encoder']
            feature_means = pd.Series(loaded['feature_means'])
            feature_columns = loaded['feature_columns']
            return model, label_encoder, feature_means, feature_columns
        except Exception:
            pass

    X, y, label_encoder, feature_means, feature_columns = load_training_data()
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, stratify=y, random_state=42
    )

    model = RandomForestClassifier(
        n_estimators=200,
        class_weight='balanced',
        random_state=42,
        n_jobs=-1,
    )
    model.fit(X_train, y_train)

    joblib.dump(
        {
            'model': model,
            'label_encoder': label_encoder,
            'feature_means': feature_means.to_dict(),
            'feature_columns': feature_columns,
        },
        MODEL_FILE,
    )

    return model, label_encoder, feature_means, feature_columns


def get_feature_description(column_name: str) -> str:
    return FEATURE_DESCRIPTIONS.get(
        column_name,
        'يُستخدم هذا العمود لمساعدة الموديل في التمييز بين حركة شبكة طبيعية وهجومية.',
    )


def call_gemini_api(text_report: str, api_key: str, system_prompt: str = DEFAULT_GEMINI_SYSTEM_PROMPT) -> str:
    """
    Call Gemini API with the text report and return the security analysis.
    Returns: The security analysis text from Gemini
    """
    if not GEMINI_AVAILABLE:
        raise RuntimeError('google.generativeai is not installed. Run: pip install google-generativeai')
    
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-pro')
        
        prompt = f"""{system_prompt}

Here is the network traffic analysis report to analyze:

{text_report}

Please provide a comprehensive security analysis report based on the data above."""
        
        response = model.generate_content(prompt, stream=False)
        return response.text
    except Exception as e:
        raise RuntimeError(f'Failed to call Gemini API: {str(e)}')


def markdown_to_docx(markdown_text: str, output_path: str) -> None:
    """
    Convert Markdown text to a professionally formatted DOCX file.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError('python-docx is not installed. Run: pip install python-docx')
    
    doc = Document()
    
    # Add title if it starts with #
    lines = markdown_text.split('\n')
    processed = 0
    
    for line in lines:
        line = line.rstrip()
        
        if not line.strip():
            doc.add_paragraph()
            processed += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            processed += 1
            continue
        elif line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=2)
            processed += 1
            continue
        elif line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=3)
            processed += 1
            continue
        
        # Handle bullet points
        if line.startswith('- '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            processed += 1
            continue
        elif line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(text, style='List Bullet')
            processed += 1
            continue
        
        # Handle numbered lists
        if re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.+)', line)
            if match:
                text = match.group(2).strip()
                p = doc.add_paragraph(text, style='List Number')
                processed += 1
                continue
        
        # Handle bold and italic
        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)  # Bold
        line = re.sub(r'\*(.+?)\*', r'\1', line)  # Italic
        
        # Default: add as normal paragraph
        doc.add_paragraph(line)
        processed += 1
    
    doc.save(output_path)


class CICIDSApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title('CICIDS2017 Intrusion Detection Desktop App')
        self.geometry('1000x720')
        self.minsize(950, 680)

        self.gemini_api_key = ''
        self.gemini_system_prompt = DEFAULT_GEMINI_SYSTEM_PROMPT
        self.processing = False

        try:
            self.model, self.label_encoder, self.feature_means, self.feature_columns = train_or_load_model()
        except Exception as ex:
            messagebox.showerror(
                'خطأ في تحميل الموديل',
                f'فشل تحميل أو تدريب الموديل: {ex}',
            )
            self.destroy()
            return

        self.top_feature_importances = self.get_top_features(3)

        self.tabview = ctk.CTkTabview(self, width=960, height=650)
        self.tabview.pack(padx=16, pady=16, expand=True, fill='both')

        self.tabview.add('Manual Prediction')
        self.tabview.add('Bulk Analysis')
        self.tabview.add('Settings')

        self.create_manual_tab(self.tabview.tab('Manual Prediction'))
        self.create_bulk_tab(self.tabview.tab('Bulk Analysis'))
        self.create_settings_tab(self.tabview.tab('Settings'))


    def get_top_features(self, n=3):
        importances = list(zip(self.feature_columns, self.model.feature_importances_))
        sorted_features = sorted(importances, key=lambda item: item[1], reverse=True)
        return sorted_features[:n]

    def create_manual_tab(self, parent):
        frame = ctk.CTkFrame(parent)
        frame.pack(expand=True, fill='both', padx=20, pady=20)

        left_frame = ctk.CTkFrame(frame)
        left_frame.pack(side='left', fill='both', expand=False, padx=(0, 20), pady=10)

        # Scrollable area to show all features (entries)
        scroll_frame = ctk.CTkScrollableFrame(left_frame, width=420, height=520)
        scroll_frame.pack(fill='both', expand=True)

        self.manual_entries = {}
        # create entries for all numeric feature columns
        for row_index, column_name in enumerate(self.feature_columns):
            label = ctk.CTkLabel(scroll_frame, text=column_name, anchor='w')
            label.grid(row=row_index, column=0, sticky='w', pady=6, padx=6)
            entry = ctk.CTkEntry(scroll_frame, width=220)
            # set default value to the column mean
            default_val = self.feature_means.get(column_name, '')
            if pd.notnull(default_val):
                entry.insert(0, str(round(float(default_val), 6)))
            entry.grid(row=row_index, column=1, pady=6, padx=6)
            self.manual_entries[column_name] = entry

        predict_button = ctk.CTkButton(
            left_frame,
            text='Predict Traffic Type',
            command=self.manual_predict,
            width=220,
        )
        predict_button.pack(pady=(12, 6))

        self.manual_output = ctk.CTkTextbox(left_frame, width=440, height=120)
        self.manual_output.pack(pady=(6, 6))
        self.manual_output.configure(state='disabled')

        info_frame = ctk.CTkFrame(frame)
        info_frame.pack(side='right', expand=True, fill='both', pady=10)

        info_title = ctk.CTkLabel(info_frame, text='Top 3 Model Features', font=ctk.CTkFont(size=16, weight='bold'))
        info_title.pack(anchor='n', pady=(10, 10))

        for feature_name, importance in self.top_feature_importances:
            desc = get_feature_description(feature_name)
            info_text = f'{feature_name}  -  importance: {importance:.4f}\n{desc}'
            item_label = ctk.CTkLabel(info_frame, text=info_text, justify='left', wraplength=320)
            item_label.pack(anchor='w', padx=10, pady=6)

    def manual_predict(self):
        try:
            row = pd.Series(self.feature_means).copy()
            for feature_name, entry in self.manual_entries.items():
                raw_value = entry.get().strip()
                if raw_value != '':
                    try:
                        row[feature_name] = float(raw_value)
                    except ValueError:
                        # keep default mean if conversion fails
                        pass

            input_vector = row.reindex(self.feature_columns).astype(float).values.reshape(1, -1)
            prediction_encoded = self.model.predict(input_vector)
            prediction_text = self.label_encoder.inverse_transform(prediction_encoded)[0]

            self.manual_output.configure(state='normal')
            self.manual_output.delete('0.0', 'end')
            self.manual_output.insert('0.0', f'Prediction: {prediction_text}')
            self.manual_output.configure(state='disabled')
        except Exception as ex:
            messagebox.showerror('Prediction Error', f'An error occurred during prediction: {ex}')

    def create_bulk_tab(self, parent):
        frame = ctk.CTkFrame(parent)
        frame.pack(expand=True, fill='both', padx=20, pady=20)

        button_frame = ctk.CTkFrame(frame)
        button_frame.pack(pady=(10, 10))

        button = ctk.CTkButton(
            button_frame,
            text='Choose Data File',
            command=self.choose_bulk_file,
            width=200,
        )
        button.pack(side='left', padx=5)

        self.gemini_checkbox = ctk.CTkCheckBox(
            button_frame,
            text='Generate AI Security Report (requires Gemini API Key)',
            onvalue=True,
            offvalue=False,
        )
        self.gemini_checkbox.pack(side='left', padx=15)

        self.bulk_info = ctk.CTkLabel(
            frame,
            text='Select a CSV or Excel file containing unlabeled network traffic data.',
            justify='center',
            wraplength=900,
        )
        self.bulk_info.pack(pady=(4, 12))

        self.bulk_textbox = ctk.CTkTextbox(frame, width=900, height=400)
        self.bulk_textbox.pack(padx=10, pady=10, expand=True, fill='both')
        self.bulk_textbox.configure(state='disabled')

    def create_settings_tab(self, parent):
        frame = ctk.CTkFrame(parent)
        frame.pack(expand=True, fill='both', padx=20, pady=20)

        # Gemini API Key section
        api_label = ctk.CTkLabel(frame, text='Gemini API Configuration', font=ctk.CTkFont(size=14, weight='bold'))
        api_label.pack(anchor='nw', pady=(10, 10))

        key_label = ctk.CTkLabel(frame, text='Gemini API Key:', anchor='w')
        key_label.pack(anchor='w', pady=(5, 2), padx=10)

        self.api_key_entry = ctk.CTkEntry(frame, width=500, show='*')
        self.api_key_entry.pack(anchor='w', pady=(2, 10), padx=10)
        if self.gemini_api_key:
            self.api_key_entry.insert(0, self.gemini_api_key)

        # System prompt section
        prompt_label = ctk.CTkLabel(frame, text='Custom System Prompt (Optional):', anchor='w')
        prompt_label.pack(anchor='w', pady=(15, 5), padx=10)

        self.system_prompt_textbox = ctk.CTkTextbox(frame, width=800, height=200)
        self.system_prompt_textbox.pack(padx=10, pady=5)
        self.system_prompt_textbox.insert('0.0', self.gemini_system_prompt)

        # Save button
        save_btn = ctk.CTkButton(frame, text='Save Settings', command=self.save_settings, width=150)
        save_btn.pack(pady=(10, 5))

        info_label = ctk.CTkLabel(
            frame,
            text='Note: Your API key is stored only in this session. Get your Gemini API key from Google AI Studio (https://aistudio.google.com/)',
            justify='left',
            wraplength=800,
            text_color='gray'
        )
        info_label.pack(anchor='w', pady=(20, 10), padx=10)

    def save_settings(self):
        api_key = self.api_key_entry.get().strip()
        system_prompt = self.system_prompt_textbox.get('0.0', 'end').strip()

        if api_key:
            self.gemini_api_key = api_key
        if system_prompt:
            self.gemini_system_prompt = system_prompt

        messagebox.showinfo('Settings Saved', 'Your settings have been saved successfully.')

    def choose_bulk_file(self):
        try:
            file_path = filedialog.askopenfilename(
                title='Select data file (CSV or Excel)',
                filetypes=[('Data Files', '*.csv *.xlsx *.xls')],
            )
            if not file_path:
                return

            self.bulk_info.configure(text=f'Processing: {os.path.basename(file_path)}')
            self.bulk_textbox.configure(state='normal')
            self.bulk_textbox.delete('0.0', 'end')
            self.bulk_textbox.insert('0.0', 'Reading file and cleaning data...')
            self.bulk_textbox.configure(state='disabled')
            self.update()

            df = self.read_input_file(file_path)
            report_path = self.generate_report(df)

            # Check if user wants Gemini analysis
            should_use_gemini = self.gemini_checkbox.get() if hasattr(self, 'gemini_checkbox') else False
            
            if should_use_gemini and self.gemini_api_key and GEMINI_AVAILABLE and DOCX_AVAILABLE:
                self.generate_gemini_and_docx(report_path)
            elif should_use_gemini and not self.gemini_api_key:
                messagebox.showinfo('Gemini Not Configured', 'Please configure your Gemini API Key in the Settings tab first.')

            messagebox.showinfo('Success', f'Statistical report saved:\n{report_path}')
            self.bulk_info.configure(text=f'Report saved: {os.path.basename(report_path)}')
        except Exception as ex:
            error_message = str(ex)
            self.bulk_textbox.configure(state='normal')
            self.bulk_textbox.delete('0.0', 'end')
            self.bulk_textbox.insert('0.0', f'Error: {error_message}')
            self.bulk_textbox.configure(state='disabled')
            messagebox.showerror('Error processing file', error_message)

    def generate_gemini_and_docx(self, txt_report_path: str):
        """
        Generate Gemini analysis and convert to DOCX with progress bar.
        """
        # Create progress window
        progress_window = ctk.CTkToplevel(self)
        progress_window.title('Generating Security Analysis Report')
        progress_window.geometry('500x200')
        progress_window.resizable(False, False)
        progress_window.attributes('-topmost', True)

        label = ctk.CTkLabel(progress_window, text='Calling Gemini API and generating report...', font=ctk.CTkFont(size=12))
        label.pack(pady=(20, 10), padx=20)

        progress_bar = ctk.CTkProgressBar(progress_window, width=450, mode='indeterminate')
        progress_bar.pack(pady=(10, 20), padx=20)
        progress_bar.start()

        status_label = ctk.CTkLabel(progress_window, text='Step 1/3: Reading statistical report...', text_color='gray')
        status_label.pack(pady=(5, 20))

        progress_window.update()

        def process_gemini():
            try:
                # Step 1: Read the text report
                status_label.configure(text='Step 1/3: Reading statistical report...')
                progress_window.update()
                time.sleep(0.5)

                with open(txt_report_path, 'r', encoding='utf-8') as f:
                    txt_content = f.read()

                # Step 2: Call Gemini
                status_label.configure(text='Step 2/3: Sending to Gemini API (this may take a moment)...')
                progress_window.update()

                gemini_response = call_gemini_api(txt_content, self.gemini_api_key, self.gemini_system_prompt)

                # Step 3: Convert to DOCX
                status_label.configure(text='Step 3/3: Converting to Word format...')
                progress_window.update()

                docx_path = filedialog.asksaveasfilename(
                    title='Save Security Analysis Report',
                    defaultextension='.docx',
                    defaultdir=os.path.dirname(txt_report_path),
                    filetypes=[('Word Documents', '*.docx')],
                )
                if not docx_path:
                    progress_window.destroy()
                    return

                markdown_to_docx(gemini_response, docx_path)

                progress_bar.stop()
                progress_window.destroy()

                messagebox.showinfo(
                    'Success',
                    f'Security analysis report created successfully!\n\n'
                    f'Saved to: {docx_path}'
                )
            except Exception as e:
                progress_bar.stop()
                progress_window.destroy()
                messagebox.showerror(
                    'Error Generating Report',
                    f'Failed to generate Gemini analysis:\n{str(e)}'
                )

        # Run Gemini call in a separate thread
        thread = threading.Thread(target=process_gemini, daemon=True)
        thread.start()

    def read_input_file(self, file_path: str) -> pd.DataFrame:
        try:
            if file_path.lower().endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path)
            else:
                df = pd.read_csv(file_path)
            return df
        except Exception as ex:
            raise ValueError(f'Failed to read file. Ensure it is a valid CSV or Excel file. Details: {ex}')

    def generate_report(self, df: pd.DataFrame) -> str:
        if df.empty:
            raise ValueError('The uploaded file is empty.')

        df_clean = df.copy()
        if 'label' in df_clean.columns:
            df_clean = df_clean.drop(columns=['label'])

        missing_columns = [col for col in self.feature_columns if col not in df_clean.columns]
        if missing_columns:
            missing_text = ', '.join(missing_columns)
            raise ValueError(f'The file does not contain the following required columns:\n{missing_text}')

        df_clean = df_clean[self.feature_columns].copy()
        df_clean.replace([np.inf, -np.inf], np.nan, inplace=True)
        df_clean.fillna(self.feature_means, inplace=True)

        try:
            X = df_clean.astype(float)
        except Exception as ex:
            raise ValueError(f'Failed to convert file data to numbers. Ensure columns contain numeric values. Details: {ex}')

        predictions_encoded = self.model.predict(X)
        predictions = self.label_encoder.inverse_transform(predictions_encoded)

        report_rows = df_clean.shape[0]
        benign_mask = np.char.lower(predictions.astype(str)) == 'benign'
        normal_indexes = np.where(benign_mask)[0].tolist()
        attack_indexes = np.where(~benign_mask)[0].tolist()

        attack_types, attack_counts = np.unique(predictions[~benign_mask], return_counts=True)
        attack_summary_list = list(zip(attack_types.tolist(), attack_counts.tolist()))

        # Compute per-attack top 5 features by combining global importances and mean-difference
        global_importances = pd.Series(self.model.feature_importances_, index=self.feature_columns)
        per_attack_top_features = {}
        X_df = X.copy()
        # ensure benign_mean exists
        if normal_indexes:
            benign_mean = X_df.iloc[normal_indexes].mean()
        else:
            benign_mean = X_df.mean()

        for attack in attack_types:
            attack_mask = predictions == attack
            if attack_mask.sum() == 0:
                per_attack_top_features[attack] = []
                continue
            attack_mean = X_df.iloc[np.where(attack_mask)[0]].mean()
            mean_diff = (attack_mean - benign_mean).abs().fillna(0)
            score = global_importances * mean_diff
            top5 = score.sort_values(ascending=False).head(5)
            per_attack_top_features[attack] = top5

        report_text = []
        report_text.append('Network Traffic Analysis Report using Random Forest')
        report_text.append('===========================================================')
        report_text.append(f'Total rows processed: {report_rows}')
        report_text.append('')
        report_text.append('Normal Traffic (BENIGN):')
        report_text.append(f'- Number of benign rows: {len(normal_indexes)}')
        report_text.append(f'- Row indexes (1-based): {", ".join(str(idx + 1) for idx in normal_indexes) if normal_indexes else "None"}')
        report_text.append('')
        report_text.append('Attacks:')
        report_text.append(f'- Number of attack rows: {len(attack_indexes)}')
        report_text.append(f'- Row indexes (1-based): {", ".join(str(idx + 1) for idx in attack_indexes) if attack_indexes else "None"}')
        report_text.append('')
        report_text.append('Attack types breakdown:')
        if attack_summary_list:
            for attack_name, count in attack_summary_list:
                report_text.append(f'- {attack_name}: {count} rows')
        else:
            report_text.append('- No attacks detected.')

        report_text.append('')
        report_text.append('Top 5 features per attack type (with short explanation):')
        for attack_name, top_series in per_attack_top_features.items():
            report_text.append(f'\nAttack: {attack_name} (top 5 features)')
            if top_series.empty:
                report_text.append('- Not enough data to compute top features for this attack.')
                continue
            for feat, score in top_series.items():
                desc = FEATURE_DESCRIPTIONS_EN.get(feat, 'Feature description not available.')
                report_text.append(f'- {feat}: score={score:.6f}. {desc}')

        report_text.append('')
        report_text.append('Notes:')
        report_text.append('Feature importance is computed from the trained Random Forest (global importance)')
        report_text.append('and combined with per-attack mean differences to highlight features most indicative for each attack in this file.')

        report_content = '\n'.join(report_text)

        save_path = filedialog.asksaveasfilename(
            title='تحديد مكان واسم حفظ ملف التقرير النصي',
            defaultextension='.txt',
            filetypes=[('Text Files', '*.txt')],
        )
        if not save_path:
            raise ValueError('No path selected for saving the report.')

        with open(save_path, 'w', encoding='utf-8') as report_file:
            report_file.write(report_content)

        self.bulk_textbox.configure(state='normal')
        self.bulk_textbox.delete('0.0', 'end')
        self.bulk_textbox.insert('0.0', report_content)
        self.bulk_textbox.configure(state='disabled')

        return save_path


if __name__ == '__main__':
    ctk.set_appearance_mode('System')
    ctk.set_default_color_theme('blue')
    app = CICIDSApp()
    app.mainloop()
