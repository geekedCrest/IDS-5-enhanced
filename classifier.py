import os
import re
import warnings
from io import BytesIO
import joblib
import numpy as np
import pandas as pd

warnings.filterwarnings('ignore', category=UserWarning)

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
MODEL_PATH  = os.path.join(BASE_DIR, 'uploaded', 'cyber_rf_model.joblib')
SCALER_PATH = os.path.join(BASE_DIR, 'uploaded', 'cyber_scaler.joblib')
ENCODER_PATH = os.path.join(BASE_DIR, 'uploaded', 'cyber_encoder.joblib')

# ── Word-report generation (Gemini AI) ──────────────────────────────────────────
# Max CSV rows classified for a single report request (keeps the web request
# responsive; the report notes when the file was larger and was truncated).
REPORT_MAX_ROWS = int(os.environ.get('REPORT_MAX_ROWS', '100000'))

GEMINI_SYSTEM_PROMPT = (
    "You are an expert Senior Cyber Security Analyst & Digital Forensics Examiner. "
    "You will receive network traffic statistics classified by a Random Forest model. "
    "Your job is to draft a formal, highly professional, and comprehensive Executive & "
    "Technical Cyber Security Report fully in English. The report must strictly include the "
    "following clearly structured sections: 1. Executive Summary (overview of network health, "
    "total rows, and attack vs. benign ratio), 2. Detailed Attack Analysis (for each detected "
    "attack type, provide a scientific explanation of how it works, and map the top 5 network "
    "features to it, explaining scientifically why these metrics spiked, such as IAT or "
    "flow_duration), 3. Risk Assessment (severity and impact on infrastructure like Web, FTP, "
    "or SSH servers), 4. Mitigation & Remediation Strategies (clear, actionable engineering "
    "solutions for each attack, e.g., Rate Limiting, Timeout adjustments, Fail2ban, and CDNs). "
    "Maintain a formal, academic, and rigorous tone, avoiding generic descriptions. "
    "Format the report using Markdown headings (#, ##, ###) and bullet lists."
)


def get_gemini_api_key():
    """Resolve the Gemini API key: env var first, then a local gitignored file."""
    key = os.environ.get('GEMINI_API_KEY', '').strip()
    if key:
        return key
    key_file = os.path.join(BASE_DIR, 'gemini_api_key.txt')
    if os.path.exists(key_file):
        with open(key_file, 'r', encoding='utf-8') as fh:
            return fh.read().strip()
    return ''

KEY_FEATURES = [
    # ── Flow identity ──────────────────────────────────────────────────────
    ('dst_port',           'Destination Port',          '80=HTTP  443=HTTPS  22=SSH  21=FTP  53=DNS  3389=RDP'),
    ('protocol',           'Protocol',                  '6 = TCP     17 = UDP     1 = ICMP     0 = Other'),
    ('flow_duration',      'Flow Duration (µs)',        'Total session length in microseconds'),
    # ── Packet counts ──────────────────────────────────────────────────────
    ('tot_fwd_pkts',       'Total Fwd Packets',         'Packets sent from source → destination'),
    ('tot_bwd_pkts',       'Total Bwd Packets',         'Packets sent from destination → source'),
    ('tot_len_fwd_pkts',   'Total Fwd Payload (bytes)', 'Total byte volume in the forward direction'),
    ('tot_len_bwd_pkts',   'Total Bwd Payload (bytes)', 'Total byte volume in the backward direction'),
    # ── Flow rates ─────────────────────────────────────────────────────────
    ('flow_byts_s',        'Flow Bytes / s',            'Total bytes transferred per second'),
    ('flow_pkts_s',        'Flow Packets / s',          'Total packets per second (both directions)'),
    ('fwd_pkts_s',         'Fwd Packets / s',           'Forward-direction packet rate'),
    ('bwd_pkts_s',         'Bwd Packets / s',           'Backward-direction packet rate'),
    # ── Packet sizes ───────────────────────────────────────────────────────
    ('fwd_pkt_len_max',    'Fwd Pkt Length Max',        'Largest forward-direction packet (bytes)'),
    ('pkt_len_mean',       'Mean Packet Length',        'Average size of all packets in the flow'),
    ('pkt_len_std',        'Packet Length Std Dev',     'Variance in packet sizes — high = mixed traffic'),
    # ── Timing ─────────────────────────────────────────────────────────────
    ('flow_iat_mean',      'Flow IAT Mean (µs)',        'Average inter-arrival time between packets'),
    ('init_fwd_win_byts',  'Init Fwd Window (bytes)',   'Initial TCP receive window size (forward)'),
    # ── TCP Flags ──────────────────────────────────────────────────────────
    ('syn_flag_cnt',       'SYN Flag Count',            'SYN flags seen — high value signals SYN flood'),
    ('ack_flag_cnt',       'ACK Flag Count',            'ACK flags — normally matches packet count'),
    ('fin_flag_cnt',       'FIN Flag Count',            'FIN flags — graceful connection teardown'),
    ('rst_flag_cnt',       'RST Flag Count',            'RST flags — forced resets, scanning indicator'),
    ('psh_flag_cnt',       'PSH Flag Count',            'PSH flags — data push events in the flow'),
    # ── Ratios ─────────────────────────────────────────────────────────────
    ('down_up_ratio',      'Down / Up Ratio',           'Bytes received ÷ bytes sent; >1 = download heavy'),
]

_scaler    = None
_encoder   = None
_model     = None
_feat_cols = None


def _load_meta():
    global _scaler, _encoder, _feat_cols
    if _scaler is not None:
        return
    _scaler    = joblib.load(SCALER_PATH)
    _encoder   = joblib.load(ENCODER_PATH)
    _feat_cols = list(_scaler.feature_names_in_)


def _load_model():
    global _model
    if _model is not None:
        return
    _load_meta()
    _model = joblib.load(MODEL_PATH, mmap_mode='r')


def get_model():
    _load_model()
    return _model


def get_model_metadata():
    _load_meta()
    return {
        'classes': list(_encoder.classes_),
        'n_features': len(_feat_cols),
        'dataset': 'CICIDS 2017 / 2018',
        'algorithm': 'Random Forest',
        'estimators': 200,
        'training_samples': 916666,
    }


def get_feature_info():
    _load_meta()
    means = dict(zip(_feat_cols, _scaler.mean_))
    info = {}
    for col, label, desc in KEY_FEATURES:
        info[col] = {
            'label': label,
            'description': desc,
            'type': 'numeric',
            'median': round(float(means.get(col, 0.0)), 4),
        }
    return info


def predict_single(row_dict):
    _load_model()
    means = dict(zip(_feat_cols, _scaler.mean_))
    row = [means.get(col, 0.0) for col in _feat_cols]
    col_idx = {col: i for i, col in enumerate(_feat_cols)}
    for key, val in row_dict.items():
        if key in col_idx and val not in (None, ''):
            try:
                row[col_idx[key]] = float(val)
            except (ValueError, TypeError):
                pass
    X = np.array(row, dtype=float).reshape(1, -1)
    X_scaled = _scaler.transform(X)
    pred_enc = _model.predict(X_scaled)
    pred_label = _encoder.inverse_transform(pred_enc)[0]
    result = {'prediction': str(pred_label)}
    if hasattr(_model, 'predict_proba'):
        probs = _model.predict_proba(X_scaled)[0]
        classes = list(_encoder.classes_)
        class_probs = sorted(zip(classes, probs), key=lambda x: x[1], reverse=True)
        result['probabilities'] = [
            {'class': c, 'prob': round(float(p) * 100, 2)}
            for c, p in class_probs[:5]
        ]
    return result


def predict_batch(rows, defaults=None):
    """Classify many feature dicts at once. `rows` is a list of {feature_name: value}.
    Returns a list of {'prediction': str, 'confidence': float} aligned with `rows`.

    `defaults` (optional) is a {feature_name: value} dict — e.g. medians from an
    uploaded CSV — used as the BASELINE for features not present in a row. Any
    feature not covered by `defaults` falls back to the scaler's training mean.
    """
    _load_model()
    if not rows:
        return []
    col_idx = {col: i for i, col in enumerate(_feat_cols)}

    base = [float(_scaler.mean_[i]) for i in range(len(_feat_cols))]
    if defaults:
        for key, val in defaults.items():
            i = col_idx.get(key)
            if i is not None and val not in (None, ''):
                try:
                    base[i] = float(val)
                except (ValueError, TypeError):
                    pass

    X = np.tile(np.array(base, dtype=float), (len(rows), 1))
    for r, row_dict in enumerate(rows):
        for key, val in row_dict.items():
            i = col_idx.get(key)
            if i is not None and val not in (None, ''):
                try:
                    X[r, i] = float(val)
                except (ValueError, TypeError):
                    pass

    X_scaled = _scaler.transform(X)
    pred_enc = _model.predict(X_scaled)
    labels = _encoder.inverse_transform(pred_enc)
    results = []
    if hasattr(_model, 'predict_proba'):
        probs = _model.predict_proba(X_scaled)
        for i, lab in enumerate(labels):
            results.append({'prediction': str(lab),
                            'confidence': round(float(probs[i].max()) * 100, 2)})
    else:
        for lab in labels:
            results.append({'prediction': str(lab), 'confidence': None})
    return results


def _norm_col(name):
    """Normalise a CSV column name to internal snake_case (e.g. 'Dst Port' → 'dst_port')."""
    import re
    return re.sub(r'[\s/\-\.]+', '_', name.strip().lower()).strip('_')


def dataset_report_from_csv(file_stream, filename=''):
    """Compute a dataset-composition report from an uploaded CSV (reads the label
    column fully to get the real per-class distribution). Mirrors the shape of the
    static /api/dataset-report payload so the Statistics panel can render it."""
    file_stream.seek(0)
    header = pd.read_csv(file_stream, nrows=0)
    cols = list(header.columns)

    label_col = None
    for c in cols:
        if c.strip().lower() in ('label', 'predicted_label', 'class', 'attack', 'category'):
            label_col = c
            break

    file_stream.seek(0)
    if label_col is None:
        # No labels — just report row/feature counts.
        first = pd.read_csv(file_stream, usecols=[cols[0]], low_memory=False, on_bad_lines='skip')
        total = int(len(first))
        return {
            'success': True, 'source': 'csv', 'has_labels': False,
            'dataset': filename or 'Uploaded CSV', 'total_samples': total,
            'n_features': len(cols), 'n_classes': 0,
            'benign_count': 0, 'benign_pct': 0.0, 'attack_count': 0, 'attack_pct': 0.0,
            'algorithm': 'Random Forest', 'estimators': 200, 'classes': [],
        }

    series = pd.read_csv(file_stream, usecols=[label_col], low_memory=False,
                         on_bad_lines='skip')[label_col].astype(str).str.strip()
    vc = series.value_counts()
    total = int(vc.sum()) or 1

    classes, benign_count = [], 0
    for name, cnt in vc.items():
        cnt = int(cnt)
        is_benign = name.lower() in ('benign', 'normal', 'benign traffic')
        if is_benign:
            benign_count += cnt
        classes.append({'name': name, 'count': cnt,
                        'pct': round(cnt / total * 100, 4), 'is_benign': is_benign})
    attack_count = total - benign_count

    return {
        'success': True, 'source': 'csv', 'has_labels': True,
        'dataset': filename or 'Uploaded CSV', 'total_samples': total,
        'n_features': max(0, len(cols) - 1), 'n_classes': len(classes),
        'benign_count': benign_count, 'benign_pct': round(benign_count / total * 100, 2),
        'attack_count': attack_count, 'attack_pct': round(attack_count / total * 100, 2),
        'algorithm': 'Random Forest', 'estimators': 200, 'classes': classes,
    }


def feature_info_from_csv(file_stream, nrows=2000):
    _load_meta()
    df = pd.read_csv(file_stream, nrows=nrows, low_memory=False, on_bad_lines='skip')

    # Drop label / metadata columns
    drop_candidates = ('Label', ' Label', 'label', 'Predicted_Label',
                       '__source_file', 'Timestamp', ' Timestamp',
                       'Src IP', 'Dst IP', 'Src Port', ' Source Port')
    df = df.drop(columns=[c for c in drop_candidates if c in df.columns])

    # Build a normalised-name → original-name lookup so we can match
    # internal names like 'dst_port' against CSV headers like 'Dst Port'
    col_map = {_norm_col(c): c for c in df.columns}

    # Some CSV variants use abbreviated names that don't normalise to the
    # model's internal names — add explicit aliases here.
    _ALIASES = {
        'tot_len_fwd_pkts': ['totlen_fwd_pkts', 'total_length_of_fwd_packets',
                             'total_fwd_packets_length', 'totlen_fwd_pkts'],
        'tot_len_bwd_pkts': ['totlen_bwd_pkts', 'total_length_of_bwd_packets',
                             'total_bwd_packets_length'],
    }
    for internal, candidates in _ALIASES.items():
        if internal not in col_map:
            for alias in candidates:
                if alias in col_map:
                    col_map[internal] = col_map[alias]
                    break

    means = dict(zip(_feat_cols, _scaler.mean_))
    info = {}
    for col, label, desc in KEY_FEATURES:
        csv_col = col_map.get(col)          # e.g. col='dst_port' → csv_col='Dst Port'
        if csv_col and pd.api.types.is_numeric_dtype(df[csv_col]):
            series = df[csv_col].dropna()
            med = float(series.median()) if not series.empty else means.get(col, 0.0)
        else:
            med = means.get(col, 0.0)
        info[col] = {
            'label': label,
            'description': desc,
            'type': 'numeric',
            'median': round(med, 4),
        }
    return info


# ════════════════════════════════════════════════════════════════════════════════
# Word-report pipeline:  CSV → bulk classify → statistics → Gemini → DOCX
# (Ports the behaviour of uploaded/cicids_app_complete.py into the web dashboard.)
# ════════════════════════════════════════════════════════════════════════════════

_ALIASES = {
    'tot_len_fwd_pkts': ['totlen_fwd_pkts', 'total_length_of_fwd_packets',
                         'total_fwd_packets_length'],
    'tot_len_bwd_pkts': ['totlen_bwd_pkts', 'total_length_of_bwd_packets',
                         'total_bwd_packets_length'],
}


def _build_feature_matrix(df):
    """Align an arbitrary CICIDS-style CSV to the model's feature columns.

    Returns a DataFrame with exactly `_feat_cols` columns (in order). Missing or
    non-numeric columns are filled with the scaler's training mean for that feature.
    Also returns the list of feature names that were actually matched from the CSV.
    """
    col_map = {_norm_col(c): c for c in df.columns}
    for internal, candidates in _ALIASES.items():
        if internal not in col_map:
            for alias in candidates:
                if alias in col_map:
                    col_map[internal] = col_map[alias]
                    break

    means = dict(zip(_feat_cols, _scaler.mean_))
    data = {}
    matched = []
    n = len(df)
    for col in _feat_cols:
        csv_col = col_map.get(col)
        if csv_col is not None:
            series = pd.to_numeric(df[csv_col], errors='coerce')
            series = series.replace([np.inf, -np.inf], np.nan)
            if series.notna().any():
                data[col] = series.fillna(means.get(col, 0.0)).to_numpy()
                matched.append(col)
                continue
        data[col] = np.full(n, means.get(col, 0.0), dtype=float)

    X = pd.DataFrame(data, columns=_feat_cols)
    return X, matched


def analyze_csv_predictions(file_stream, max_rows=REPORT_MAX_ROWS):
    """Read a CSV, classify every row, and return (predictions, X, matched, total, truncated)."""
    _load_model()
    # Read one extra row to detect truncation
    df = pd.read_csv(file_stream, nrows=max_rows + 1, low_memory=False, on_bad_lines='skip')
    truncated = len(df) > max_rows
    if truncated:
        df = df.iloc[:max_rows]
    if df.empty:
        raise ValueError('The uploaded CSV contains no data rows.')

    # Drop obvious label / metadata columns so they don't pollute feature matching
    drop_candidates = ('Label', ' Label', 'label', 'Predicted_Label',
                       '__source_file', 'Timestamp', ' Timestamp',
                       'Src IP', 'Dst IP', 'Flow ID')
    df = df.drop(columns=[c for c in drop_candidates if c in df.columns])

    X, matched = _build_feature_matrix(df)
    if not matched:
        raise ValueError('No recognisable CICIDS feature columns were found in this CSV.')

    X_scaled = _scaler.transform(X)
    pred_enc = _model.predict(X_scaled)
    preds = _encoder.inverse_transform(pred_enc)
    return preds, X, matched, len(df), truncated


def _is_benign(label):
    return str(label).strip().lower() in ('benign', 'normal', 'benign traffic')


def build_statistics_text(preds, X, matched, total_rows, truncated):
    """Build the plain-text statistical report fed to Gemini (ports generate_statistics)."""
    preds = np.asarray(preds)
    unique_classes, counts = np.unique(preds, return_counts=True)
    class_counts = dict(zip(unique_classes.tolist(), counts.tolist()))

    benign_count = sum(c for k, c in class_counts.items() if _is_benign(k))
    attack_count = total_rows - benign_count

    benign_idx = [i for i, p in enumerate(preds) if _is_benign(p)]
    attack_idx = [i for i, p in enumerate(preds) if not _is_benign(p)]

    importances = pd.Series(_model.feature_importances_, index=_feat_cols)
    top_features = importances.nlargest(5)

    lines = [
        '=' * 60,
        'NETWORK TRAFFIC ANALYSIS REPORT',
        'Random Forest Classification Results (CICIDS 2017/2018)',
        '=' * 60,
        '',
        f'Total Rows Analyzed: {total_rows}',
        f'Benign Traffic: {benign_count} ({benign_count / total_rows * 100:.1f}%)',
        f'Attack Traffic: {attack_count} ({attack_count / total_rows * 100:.1f}%)',
        f'Feature Columns Matched From CSV: {len(matched)} / {len(_feat_cols)}',
    ]
    if truncated:
        lines.append(f'NOTE: File was larger than the {total_rows}-row limit; only the '
                     f'first {total_rows} rows were classified for this report.')
    lines += [
        '',
        'Benign Row Indexes (1-based, first 50):',
        f'{", ".join(str(i + 1) for i in benign_idx[:50])}{"..." if len(benign_idx) > 50 else ""}',
        '',
        'Attack Row Indexes (1-based, first 50):',
        f'{", ".join(str(i + 1) for i in attack_idx[:50])}{"..." if len(attack_idx) > 50 else ""}',
        '',
        'Attack Type Distribution:',
    ]
    attack_items = sorted(((k, c) for k, c in class_counts.items() if not _is_benign(k)),
                          key=lambda kv: kv[1], reverse=True)
    if attack_items:
        for attack_type, count in attack_items:
            lines.append(f'  {attack_type}: {count} rows ({count / total_rows * 100:.1f}%)')
    else:
        lines.append('  No attacks detected.')

    lines += ['', 'Top 5 Most Important Features (global model importance):']
    for feat, importance in top_features.items():
        lines.append(f'  {feat}: {importance:.4f}')

    def _mean(col):
        return float(X[col].mean()) if col in X.columns else 0.0

    lines += [
        '',
        'Statistical Summary (means over analyzed rows):',
        f'  Mean flow duration: {_mean("flow_duration"):.2f}',
        f'  Mean forward packets: {_mean("tot_fwd_pkts"):.2f}',
        f'  Mean backward packets: {_mean("tot_bwd_pkts"):.2f}',
        f'  Mean bytes/sec: {_mean("flow_byts_s"):.2f}',
        '',
    ]

    stats = {
        'total_rows': total_rows,
        'benign_count': benign_count,
        'attack_count': attack_count,
        'attack_distribution': dict(attack_items),
        'matched_features': len(matched),
        'truncated': truncated,
    }
    return '\n'.join(lines), stats


def call_gemini_api(statistical_report):
    """Send the statistical report to Gemini and return the Markdown security analysis."""
    api_key = get_gemini_api_key()
    if not api_key:
        raise RuntimeError('No Gemini API key configured. Set GEMINI_API_KEY or create '
                           'gemini_api_key.txt in the project root.')
    try:
        from google import genai
        from google.genai import types
    except ImportError:
        raise RuntimeError('google-genai not installed. Run: pip install google-genai')

    client = genai.Client(api_key=api_key)
    prompt = (
        f"{GEMINI_SYSTEM_PROMPT}\n\n"
        "====================================\n"
        "Network traffic analysis statistics:\n"
        "====================================\n\n"
        f"{statistical_report}\n\n"
        "====================================\n"
        "Now, write the comprehensive security report:\n"
        "===================================="
    )
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
        config=types.GenerateContentConfig(
            system_instruction=GEMINI_SYSTEM_PROMPT,
            temperature=0.2,
        ),
    )
    if not response.text:
        raise RuntimeError('Gemini returned an empty response.')
    return response.text


def _add_inline_md(paragraph, text):
    """Add `text` to a docx paragraph, rendering **bold** / *italic* markdown inline."""
    # Split on **bold** and *italic* while keeping the delimiters
    for part in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def markdown_to_docx(markdown_text, output):
    """Convert Markdown text to a formatted DOCX. `output` may be a path or file-like object."""
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()
    doc.add_heading('Digital Forensics & Network Traffic Analysis Report (CICIDS2017)', level=0)

    for raw in markdown_text.split('\n'):
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(('- ', '* ')):
            _add_inline_md(doc.add_paragraph(style='List Bullet'), line[2:].strip())
        elif re.match(r'^\d+\.\s', line):
            m = re.match(r'^(\d+)\.\s(.+)', line)
            _add_inline_md(doc.add_paragraph(style='List Number'), m.group(2).strip() if m else line)
        else:
            _add_inline_md(doc.add_paragraph(), line)

    doc.save(output)


def generate_security_report(file_stream, max_rows=REPORT_MAX_ROWS):
    """Full pipeline: CSV → classify → statistics → Gemini → DOCX bytes.

    Returns (BytesIO docx, meta dict). Raises on failure (no key, Gemini error, etc.).
    """
    preds, X, matched, total, truncated = analyze_csv_predictions(file_stream, max_rows=max_rows)
    stats_text, stats = build_statistics_text(preds, X, matched, total, truncated)
    gemini_md = call_gemini_api(stats_text)

    bio = BytesIO()
    markdown_to_docx(gemini_md, bio)
    bio.seek(0)
    return bio, stats
